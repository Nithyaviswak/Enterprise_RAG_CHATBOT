"""Document ingestion: page-aware parsing, metadata, and duplicate detection.

Extracts text from PDF / DOCX / TXT / MD / CSV while preserving provenance
metadata (source filename, page number, chunk id, document id). Prevents
duplicate ingestion via a content hash that is compared against the vector
store.
"""

import hashlib
import logging
import os
from dataclasses import dataclass, field
from typing import Optional

from app.rag.failures import InvalidDocumentError, ParsingFailure
from app.rag.chunking import chunk_by_words

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv"}

_TXT_LIKE = {".txt", ".md", ".csv"}


@dataclass
class Page:
    """A page (or logical section) of a parsed document."""

    text: str
    page_number: Optional[int]  # None for single-page text formats


@dataclass
class ParsedDocument:
    """Result of parsing a document into pages."""

    filename: str
    file_type: str
    pages: list[Page] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text)

    @property
    def word_count(self) -> int:
        return len(self.full_text.split())


def is_allowed(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in ALLOWED_EXTENSIONS


def content_hash(text: str) -> str:
    """Return a stable sha256 hash of normalized text for dedup checks."""
    normalized = " ".join(text.split()).lower()
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def parse_document(file_path: str, file_type: str, filename: Optional[str] = None) -> ParsedDocument:
    """Parse a file into page-aware text units.

    Metadata safety: raises :class:`ParsingFailure` / :class:`InvalidDocumentError`
    with an explicit failure category instead of leaking raw exceptions.
    """
    ext = file_type.lower() if file_type.startswith(".") else f".{file_type.lower()}"
    filename = filename or os.path.basename(file_path)
    ext = ext.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise InvalidDocumentError(
            f"Unsupported file type: {ext}. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
            filename=filename,
        )

    if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
        raise InvalidDocumentError("File is missing or empty.", filename=filename)

    try:
        if ext == ".pdf":
            return _parse_pdf(file_path, filename)
        if ext == ".docx":
            return _parse_docx(file_path, filename)
        return _parse_txt(file_path, filename, ext)
    except InvalidDocumentError:
        raise
    except Exception as e:  # pragma: no cover - defensive
        raise ParsingFailure(f"Failed to parse document: {e}", filename=filename)


def _parse_pdf(file_path: str, filename: str) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages: list[Page] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as e:  # pragma: no cover
            logger.warning("Could not extract text from %s page %d: %s", filename, i, e)
            text = ""
        if text:
            pages.append(Page(text=text, page_number=i))
    if not pages:
        raise ParsingFailure("PDF contains no extractable text.", filename=filename)
    return ParsedDocument(filename=filename, file_type=".pdf", pages=pages)


def _parse_docx(file_path: str, filename: str) -> ParsedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    texts: list[str] = []
    buffer: list[str] = []

    def flush():
        if buffer:
            text = " ".join(t for t in buffer if t).strip()
            if text:
                texts.append(text)
            buffer.clear()

    for para in doc.paragraphs:
        if not para.text.strip():
            flush()  # paragraph break ≈ section boundary
        else:
            buffer.append(para.text.strip())
    flush()

    if not texts:
        raise InvalidDocumentError("DOCX contains no extractable text.", filename=filename)
    pages = [Page(text=t, page_number=None) for t in texts]
    return ParsedDocument(filename=filename, file_type=".docx", pages=pages)


def _parse_txt(file_path: str, filename: str, ext: str) -> ParsedDocument:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    if not text.strip():
        raise InvalidDocumentError("File contains no extractable text.", filename=filename)
    return ParsedDocument(filename=filename, file_type=ext, pages=[Page(text=text.strip(), page_number=None)])


def build_chunks(
    parsed: ParsedDocument,
    document_id: str,
    chunk_size: Optional[int] = None,
    overlap: Optional[int] = None,
) -> list[dict]:
    """Chunk a parsed document and attach provenance metadata to each chunk.

    Returns a list of dicts:
        {id, content, metadata}
    where metadata contains: source (filename), document_id, chunk_id,
    chunk_index, file_type, page (page number or None), content_hash.
    """
    chunks_out: list[dict] = []
    for page in parsed.pages:
        page_chunks = chunk_by_words(page.text, chunk_size=chunk_size, overlap=overlap)
        for idx, chunk_text in enumerate(page_chunks):
            chunk_id = f"{document_id}_{len(chunks_out)}"
            chunks_out.append(
                {
                    "id": chunk_id,
                    "content": chunk_text,
                    "metadata": {
                        "source": parsed.filename,
                        "document_id": document_id,
                        "chunk_id": chunk_id,
                        "chunk_index": len(chunks_out),
                        "file_type": parsed.file_type,
                        "page": page.page_number,
                        "content_hash": content_hash(chunk_text),
                    },
                }
            )
    return chunks_out


def is_duplicate_in_store(collection, document_id: Optional[str], content_hash_value: str) -> bool:
    """Return True when a chunk with the same content hash already exists.

    Prevents re-ingesting the same file (or identical content) multiple times.
    """
    try:
        results = collection.get(where={"content_hash": content_hash_value}, limit=1)
        return bool(results and results.get("ids"))
    except Exception as e:
        logger.warning("Duplicate check failed: %s", e)
        return False